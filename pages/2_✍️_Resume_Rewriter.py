"""
Resume Rewriter Page
AI-powered resume optimization and rewriting suggestions
"""

import streamlit as st
import sys
from pathlib import Path
from datetime import datetime
import json
import re
from io import BytesIO
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.config import Config, COLORS
from src.utils.file_parser import parse_file
from src.utils.helpers import get_file_size_mb, validate_file_extension
from src.core.resume_analyzer import ResumeAnalyzer
from src.core.resume_rewriter import ResumeRewriter

st.set_page_config(page_title="Resume Rewriter", page_icon="✍️", layout="wide")

st.markdown(f"""
    <style>
        .improvement-card {{
            background-color: #f8f9fa;
            border-left: 4px solid {COLORS['warning']};
            padding: 1rem;
            margin: 0.5rem 0;
            border-radius: 5px;
        }}
        .priority-high {{
            background-color: #fff3cd;
            border-left: 4px solid {COLORS['danger']};
            padding: 1rem;
            margin: 0.5rem 0;
            border-radius: 5px;
        }}
        .priority-medium {{
            background-color: #d1ecf1;
            border-left: 4px solid {COLORS['info']};
            padding: 1rem;
            margin: 0.5rem 0;
            border-radius: 5px;
        }}
        .example-box {{
            background-color: #e7f3ff;
            padding: 1rem;
            border-radius: 5px;
            margin: 0.5rem 0;
        }}
    </style>
""", unsafe_allow_html=True)


def extract_candidate_profile(resume_text: str, analysis: dict) -> dict:
    """
    Extract candidate profile from resume using resume analyzer
    
    Args:
        resume_text: The resume text
        analysis: Analysis from ResumeAnalyzer
        
    Returns:
        dict: Candidate profile information
    """
    profile = {
        "name": analysis.get("contact_info", {}).get("name", "Candidate"),
        "email": analysis.get("contact_info", {}).get("email", ""),
        "phone": analysis.get("contact_info", {}).get("phone", ""),
        "experience_years": analysis.get("experience_years", 0),
        "education_level": analysis.get("education_level", "Not specified"),
        "technical_skills": analysis.get("technical_skills", []),
        "soft_skills": analysis.get("soft_skills", []),
        "sections_found": analysis.get("sections_found", []),
        "current_ats_score": analysis.get("ats_score", 0),
        "strengths": analysis.get("strengths", []),
        "weaknesses": analysis.get("weaknesses", [])
    }
    return profile


def analyze_weak_phrases_with_ai(client, resume_text: str, profile: dict) -> list:
    """
    Use AI to analyze weak phrases in the resume
    
    Args:
        client: OpenAI client
        resume_text: The resume text
        profile: Candidate profile
        
    Returns:
        list: Weak phrases with alternatives
    """
    prompt = f"""You are an expert resume writer. Analyze this resume for weak/passive phrases.

CANDIDATE PROFILE:
- Experience: {profile['experience_years']} years
- Education: {profile['education_level']}
- Key Skills: {', '.join(profile['technical_skills'][:5])}

RESUME TEXT:
{resume_text}

Identify 5-10 SPECIFIC weak phrases from THIS resume and provide:
1. Original phrase (exact quote)
2. Why it's weak
3. Stronger alternative (contextualized)
4. Priority (High/Medium/Low)

Return as JSON array:
[
  {{"original": "...", "reason": "...", "alternative": "...", "priority": "High"}},
  ...
]"""

    try:
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "https://github.com/FEDI-HASSINE/cv-ai-",
                "X-Title": "UtopiaHire Resume Rewriter",
            },
            model="deepseek/deepseek-r1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        
        response_text = completion.choices[0].message.content
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        else:
            json_text = response_text
        
        return json.loads(json_text)
    except Exception as e:
        st.warning(f"⚠️ Weak phrases analysis failed: {e}")
        return []


def suggest_action_verbs_with_ai(client, resume_text: str, profile: dict) -> dict:
    """
    Use AI to suggest contextualized action verbs
    
    Args:
        client: OpenAI client
        resume_text: The resume text
        profile: Candidate profile
        
    Returns:
        dict: Action verbs by category with examples
    """
    prompt = f"""You are an expert resume writer. Suggest powerful action verbs for this candidate.

CANDIDATE PROFILE:
- Experience: {profile['experience_years']} years
- Education: {profile['education_level']}
- Key Skills: {', '.join(profile['technical_skills'][:5])}
- Sections: {', '.join(profile['sections_found'])}

RESUME TEXT (excerpt):
{resume_text[:1000]}...

Suggest 10-15 action verbs categorized by:
- Leadership (if applicable)
- Technical/Innovation
- Achievement
- Collaboration

For EACH verb, provide a contextualized example using this candidate's background.

Return as JSON:
{{
  "Leadership": [{{"verb": "...", "example": "Led team of X..."}}, ...],
  "Technical": [{{"verb": "...", "example": "Developed..."}}, ...],
  ...
}}"""

    try:
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "https://github.com/FEDI-HASSINE/cv-ai-",
                "X-Title": "UtopiaHire Resume Rewriter",
            },
            model="deepseek/deepseek-r1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        
        response_text = completion.choices[0].message.content
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        else:
            json_text = response_text
        
        return json.loads(json_text)
    except Exception as e:
        st.warning(f"⚠️ Action verbs suggestion failed: {e}")
        return {}


def suggest_quantifications_with_ai(client, resume_text: str, profile: dict) -> list:
    """
    Use AI to suggest quantification opportunities
    
    Args:
        client: OpenAI client
        resume_text: The resume text
        profile: Candidate profile
        
    Returns:
        list: Quantification opportunities
    """
    prompt = f"""You are an expert resume writer. Identify quantification opportunities in this resume.

CANDIDATE PROFILE:
- Experience: {profile['experience_years']} years
- Field: {', '.join(profile['technical_skills'][:3])}

RESUME TEXT:
{resume_text}

Find 5-10 statements that need quantifiable metrics. For each:
1. Original vague statement (exact quote)
2. Quantified version with example metrics
3. Impact explanation

Return as JSON array:
[
  {{"original": "...", "quantified": "...", "impact": "..."}},
  ...
]"""

    try:
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "https://github.com/FEDI-HASSINE/cv-ai-",
                "X-Title": "UtopiaHire Resume Rewriter",
            },
            model="deepseek/deepseek-r1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        
        response_text = completion.choices[0].message.content
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        else:
            json_text = response_text
        
        return json.loads(json_text)
    except Exception as e:
        st.warning(f"⚠️ Quantification analysis failed: {e}")
        return []


def suggest_formatting_with_ai(client, resume_text: str, profile: dict) -> list:
    """
    Use AI to suggest formatting improvements
    
    Args:
        client: OpenAI client
        resume_text: The resume text
        profile: Candidate profile
        
    Returns:
        list: Formatting suggestions
    """
    prompt = f"""You are an expert resume writer. Analyze formatting issues in this resume.

CANDIDATE PROFILE:
- Sections present: {', '.join(profile['sections_found'])}
- ATS Score: {profile['current_ats_score']}/100

RESUME TEXT:
{resume_text}

List 5-8 SPECIFIC formatting improvements needed for THIS resume.
Be actionable and specific.

Return as JSON array:
["Suggestion 1", "Suggestion 2", ...]"""

    try:
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "https://github.com/FEDI-HASSINE/cv-ai-",
                "X-Title": "UtopiaHire Resume Rewriter",
            },
            model="deepseek/deepseek-r1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1500
        )
        
        response_text = completion.choices[0].message.content
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        else:
            json_text = response_text
        
        return json.loads(json_text)
    except Exception as e:
        st.warning(f"⚠️ Formatting analysis failed: {e}")
        return []


def generate_optimized_resume_content(client, resume_text: str, profile: dict, improvements: dict) -> dict:
    """
    Generate complete optimized resume content with AI
    
    Args:
        client: OpenAI client
        resume_text: Original resume text
        profile: Candidate profile
        improvements: Suggested improvements
        
    Returns:
        dict: Optimized resume sections
    """
    prompt = f"""You are an expert resume writer. Create an optimized version of this resume.

CANDIDATE PROFILE:
- Name: {profile['name']}
- Experience: {profile['experience_years']} years
- Education: {profile['education_level']}
- Technical Skills: {', '.join(profile['technical_skills'])}
- Soft Skills: {', '.join(profile['soft_skills'])}

ORIGINAL RESUME:
{resume_text}

KEY IMPROVEMENTS TO APPLY:
- Weak phrases to fix: {len(improvements.get('weak_phrases', []))}
- Add quantifiable metrics
- Use stronger action verbs
- Improve formatting

CREATE OPTIMIZED RESUME with these sections:
1. Professional Summary (3-4 lines, compelling)
2. Professional Experience (bullet points, quantified achievements)
3. Skills (technical and soft skills)
4. Education

Maintain factual accuracy - only optimize language, don't invent information.

Return as JSON:
{{
  "summary": "...",
  "experience": ["• Achievement 1...", "• Achievement 2...", ...],
  "skills": "Technical Skills: ...\\nSoft Skills: ...",
  "education": "..."
}}"""

    try:
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": "https://github.com/FEDI-HASSINE/cv-ai-",
                "X-Title": "UtopiaHire Resume Rewriter",
            },
            model="deepseek/deepseek-r1",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=3000
        )
        
        response_text = completion.choices[0].message.content
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        else:
            json_text = response_text
        
        return json.loads(json_text)
    except Exception as e:
        st.warning(f"⚠️ Resume generation failed: {e}")
        return {}


def optimize_with_deepseek(resume_text: str, basic_analysis: dict) -> dict:
    """
    Optimisation avancée du CV avec DeepSeek R1 - Version modulaire
    
    Args:
        resume_text: Texte complet du CV
        basic_analysis: Analyse basique existante (fallback)
        
    Returns:
        dict: Suggestions optimisées avec DeepSeek ou fallback
    """
    if not st.session_state.get('use_rewriter_deepseek', False):
        # Fallback : retourner l'analyse basique
        return basic_analysis
    
    try:
        client = st.session_state['rewriter_deepseek_client']
        
        # Step 1: Extract candidate profile from resume analyzer
        st.info("📋 Step 1/5: Extracting candidate profile...")
        profile = extract_candidate_profile(resume_text, basic_analysis)
        
        # Step 2: Analyze weak phrases with context
        st.info("🔍 Step 2/5: Analyzing weak phrases...")
        weak_phrases = analyze_weak_phrases_with_ai(client, resume_text, profile)
        
        # Step 3: Suggest action verbs with context
        st.info("💪 Step 3/5: Suggesting action verbs...")
        action_verbs = suggest_action_verbs_with_ai(client, resume_text, profile)
        
        # Step 4: Suggest quantifications with context
        st.info("📊 Step 4/5: Finding quantification opportunities...")
        quantifications = suggest_quantifications_with_ai(client, resume_text, profile)
        
        # Step 5: Suggest formatting improvements
        st.info("🎨 Step 5/5: Analyzing formatting...")
        formatting = suggest_formatting_with_ai(client, resume_text, profile)
        
        # Compile all improvements
        improvements = {
            'weak_phrases': weak_phrases,
            'action_verbs': action_verbs,
            'quantifications': quantifications,
            'formatting': formatting
        }
        
        # Calculate scores
        current_score = basic_analysis.get('overall_score', basic_analysis.get('ats_score', 50))
        potential_score = min(100, current_score + 20 + len(weak_phrases) * 2)
        
        # Generate priority recommendations
        priority_recommendations = {
            'high_priority': [],
            'medium_priority': [],
            'low_priority': []
        }
        
        if weak_phrases:
            priority_recommendations['high_priority'].append(
                f"Replace {len(weak_phrases)} weak phrases with stronger alternatives"
            )
        if quantifications:
            priority_recommendations['high_priority'].append(
                f"Add quantifiable metrics to {len(quantifications)} achievements"
            )
        if action_verbs:
            priority_recommendations['medium_priority'].append(
                "Diversify action verbs across different categories"
            )
        if formatting:
            priority_recommendations['medium_priority'].append(
                f"Address {len(formatting)} formatting issues"
            )
        
        # Enrichir l'analyse basique avec DeepSeek
        enhanced_analysis = basic_analysis.copy()
        enhanced_analysis.update({
            'deepseek_enabled': True,
            'deepseek_weak_phrases': weak_phrases,
            'deepseek_action_verbs': action_verbs,
            'deepseek_quantification': quantifications,
            'deepseek_formatting': formatting,
            'deepseek_recommendations': priority_recommendations,
            'deepseek_optimized': {},  # Will be generated separately when user wants to download
            'deepseek_scores': {
                'current': current_score,
                'potential': potential_score,
                'areas': ['Language', 'Quantification', 'Formatting']
            },
            'candidate_profile': profile  # Store profile for later use
        })
        
        st.success("✅ Analyse DeepSeek R1 terminée avec succès!")
        return enhanced_analysis
        
    except Exception as e:
        st.error(f"❌ Erreur DeepSeek R1: {str(e)}")
        st.info("💡 Utilisation de l'analyse standard")
        return basic_analysis


def generate_optimized_resume_docx(original_text: str, optimized_data: dict) -> BytesIO:
    """
    Génère un CV optimisé au format DOCX
    
    Args:
        original_text: Texte original du CV
        optimized_data: Données d'optimisation (DeepSeek ou standard)
        
    Returns:
        BytesIO: Document DOCX en mémoire
    """
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Titre
    title = doc.add_heading('OPTIMIZED RESUME', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Note d'optimisation
    note = doc.add_paragraph()
    note.add_run('✨ This resume has been optimized using AI-powered suggestions\n').bold = True
    note.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Si DeepSeek a optimisé des sections
    if optimized_data.get('deepseek_enabled') and optimized_data.get('deepseek_optimized'):
        optimized_sections = optimized_data['deepseek_optimized']
        
        # Summary optimisé
        if optimized_sections.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', 1)
            doc.add_paragraph(optimized_sections['summary'])
            doc.add_paragraph()
        
        # Experience optimisée
        if optimized_sections.get('experience'):
            doc.add_heading('PROFESSIONAL EXPERIENCE', 1)
            experience_data = optimized_sections['experience']
            # Handle different data types
            if isinstance(experience_data, str):
                doc.add_paragraph(experience_data)
            elif isinstance(experience_data, list):
                for exp in experience_data:
                    if isinstance(exp, dict):
                        p = doc.add_paragraph()
                        p.add_run(exp.get('title', '')).bold = True
                        doc.add_paragraph(exp.get('description', ''))
                    elif isinstance(exp, str):
                        doc.add_paragraph(exp)
            doc.add_paragraph()
        
        # Skills
        if optimized_sections.get('skills'):
            doc.add_heading('SKILLS', 1)
            doc.add_paragraph(optimized_sections['skills'])
            doc.add_paragraph()
    else:
        # Fallback : utiliser le texte original avec suggestions
        doc.add_heading('ORIGINAL CONTENT', 1)
        doc.add_paragraph(original_text)
        doc.add_paragraph()
    
    # Section des suggestions appliquées
    doc.add_page_break()
    doc.add_heading('OPTIMIZATION NOTES', 1)
    
    if optimized_data.get('deepseek_recommendations'):
        recs = optimized_data['deepseek_recommendations']
        
        if recs.get('high_priority'):
            doc.add_heading('High Priority Changes Applied:', 2)
            high_priority = recs['high_priority']
            if isinstance(high_priority, list):
                for rec in high_priority:
                    rec_text = rec if isinstance(rec, str) else str(rec)
                    doc.add_paragraph(f"• {rec_text}", style='List Bullet')
            doc.add_paragraph()
        
        if recs.get('medium_priority'):
            doc.add_heading('Medium Priority Improvements:', 2)
            medium_priority = recs['medium_priority']
            if isinstance(medium_priority, list):
                for rec in medium_priority:
                    rec_text = rec if isinstance(rec, str) else str(rec)
                    doc.add_paragraph(f"• {rec_text}", style='List Bullet')
            doc.add_paragraph()
    
    # Sauvegarder en mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_optimized_resume_pdf(original_text: str, optimized_data: dict) -> BytesIO:
    """
    Génère un CV optimisé au format PDF
    
    Args:
        original_text: Texte original du CV
        optimized_data: Données d'optimisation
        
    Returns:
        BytesIO: Document PDF en mémoire
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                            topMargin=1*inch, bottomMargin=1*inch)
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#2E86AB',
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='#2E86AB',
        spaceAfter=6
    )
    
    # Contenu
    story = []
    
    # Titre
    story.append(Paragraph("OPTIMIZED RESUME", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Note
    note_text = f"✨ <b>This resume has been optimized using AI-powered suggestions</b><br/>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    story.append(Paragraph(note_text, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Contenu optimisé
    if optimized_data.get('deepseek_enabled') and optimized_data.get('deepseek_optimized'):
        optimized_sections = optimized_data['deepseek_optimized']
        
        if optimized_sections.get('summary'):
            story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
            story.append(Paragraph(optimized_sections['summary'], styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Autres sections...
    else:
        story.append(Paragraph("ORIGINAL CONTENT", heading_style))
        story.append(Paragraph(original_text.replace('\n', '<br/>'), styles['Normal']))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


def main():
    st.title("✍️ Resume Rewriter")
    st.markdown("Get AI-powered suggestions to optimize your resume")
    
    # Initialize
    if 'rewriter' not in st.session_state:
        st.session_state.rewriter = ResumeRewriter()
    if 'analyzer' not in st.session_state:
        st.session_state.analyzer = ResumeAnalyzer()
    
    config = Config()
    
    # Sidebar
    with st.sidebar:
        st.markdown("### ✨ Rewriting Features")
        st.markdown("""
        - **Weak Phrase Detection**
        - **Action Verb Suggestions**
        - **Quantification Opportunities**
        - **Formatting Improvements**
        - **Content Enhancement**
        - **Optimized Examples**
        """)
        
        st.markdown("---")
        
        # AI Configuration
        st.markdown("### 🤖 AI Configuration")
        st.markdown("*Pour des suggestions personnalisées*")
        
        ai_model = st.radio(
            "Moteur d'optimisation",
            options=["DeepSeek R1 (Recommandé)", "Analyse Standard"],
            index=0,
            help="DeepSeek R1 génère des suggestions personnalisées basées sur votre CV"
        )
        
        if ai_model == "DeepSeek R1 (Recommandé)":
            st.markdown("#### 🚀 DeepSeek R1 Configuration")
            
            openrouter_key = st.text_input(
                "OpenRouter API Key",
                type="password",
                help="Clé API OpenRouter pour DeepSeek R1",
                placeholder="sk-or-v1-...",
                key="rewriter_openrouter_key"
            )
            
            if openrouter_key:
                st.success("✅ DeepSeek R1 activé")
                st.session_state['rewriter_deepseek_client'] = OpenAI(
                    base_url="https://openrouter.ai/api/v1",
                    api_key=openrouter_key,
                )
                st.session_state['use_rewriter_deepseek'] = True
            else:
                st.warning("⚠️ Clé API requise pour suggestions personnalisées")
                st.session_state['use_rewriter_deepseek'] = False
                st.info("💡 Sans clé : suggestions basiques maintenues")
        else:
            st.session_state['use_rewriter_deepseek'] = False
            st.info("💡 Mode standard : suggestions basiques")
        
        st.markdown("---")
        st.markdown("### 💡 Best Practices")
        st.info("""
        - Use strong action verbs
        - Quantify achievements
        - Focus on results
        - Keep it concise
        - Tailor to job descriptions
        """)
    
    # Main content
    st.markdown("---")
    
    # Upload
    uploaded_file = st.file_uploader(
        "Upload your resume for rewriting suggestions",
        type=['pdf', 'docx', 'doc', 'txt'],
        help="We'll analyze your content and provide optimization suggestions"
    )
    
    if uploaded_file is not None:
        file_bytes = uploaded_file.read()
        file_size = get_file_size_mb(file_bytes)
        
        if file_size > config.MAX_FILE_SIZE_MB:
            st.error(f"File too large: {file_size:.2f} MB (max: {config.MAX_FILE_SIZE_MB} MB)")
            return
        
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        # Parse
        with st.spinner("📖 Parsing resume..."):
            parsed = parse_file(uploaded_file.name, file_bytes)
        
        if parsed["status"] == "error":
            st.error(f"Error: {parsed['error']}")
            return
        
        # Options
        col1, col2 = st.columns(2)
        
        with col1:
            include_analysis = st.checkbox(
                "Include resume analysis",
                value=True,
                help="Analyze resume first for better suggestions"
            )
        
        with col2:
            show_examples = st.checkbox(
                "Show optimized examples",
                value=True,
                help="Display example sections"
            )
        
        # Analyze & Rewrite button
        if st.button("✨ Generate Rewriting Suggestions", type="primary"):
            with st.spinner("🤖 Analyzing and generating suggestions..."):
                # Optional analysis
                analysis = None
                if include_analysis:
                    analysis = st.session_state.analyzer.analyze(parsed["text"])
                    st.session_state.analysis = analysis
                
                # Check if we're using AI mode (DeepSeek)
                use_ai_mode = st.session_state.get('use_rewriter_deepseek', False)
                
                # Rewrite suggestions (only use static objects in fallback mode)
                rewrite = st.session_state.rewriter.rewrite(
                    parsed["text"], 
                    analysis,
                    use_static_objects=not use_ai_mode  # Disable static objects when using AI
                )
                
                # Enrichir avec DeepSeek si activé
                if st.session_state.get('use_rewriter_deepseek', False):
                    rewrite = optimize_with_deepseek(parsed["text"], rewrite)
                
                st.session_state.rewrite = rewrite
        
        # Display results
        if 'rewrite' in st.session_state:
            rewrite = st.session_state.rewrite
            
            st.markdown("---")
            st.markdown("## ✨ Optimization Suggestions")
            
            # Afficher le score DeepSeek si disponible
            if rewrite.get('deepseek_enabled') and rewrite.get('deepseek_scores'):
                scores = rewrite['deepseek_scores']
                col1, col2 = st.columns(2)
                
                with col1:
                    st.metric(
                        "📊 Current Score",
                        f"{scores.get('current', 50)}/100",
                        help="Score actuel de votre CV"
                    )
                
                with col2:
                    improvement = scores.get('potential', 80) - scores.get('current', 50)
                    st.metric(
                        "🚀 Potential Score",
                        f"{scores.get('potential', 80)}/100",
                        delta=f"+{improvement}",
                        help="Score après optimisations"
                    )
                
                st.progress(scores.get('potential', 80) / 100)
            else:
                # Fallback : affichage standard
                improvement_score = rewrite.get('improvement_score', 50)
                st.markdown(f"""
                <div class="improvement-card">
                    <h3>📈 Improvement Potential: {improvement_score}/100</h3>
                    <p>Your resume has significant potential for optimization!</p>
                </div>
                """, unsafe_allow_html=True)
                st.progress(improvement_score / 100)
            
            # Tabs
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "🔄 Weak Phrases",
                "💪 Action Verbs",
                "📊 Quantification",
                "🎨 Formatting",
                "📝 Examples"
            ])
            
            with tab1:
                st.markdown("### 🔄 Weak Phrases to Improve")
                
                if rewrite['weak_phrases']:
                    st.info(f"Found {len(rewrite['weak_phrases'])} weak phrases that can be strengthened")
                    
                    for idx, wp in enumerate(rewrite['weak_phrases'][:10], 1):
                        st.markdown(f"""
                        <div class="improvement-card">
                            <strong>Example {idx}</strong><br>
                            &#10060; <span style="color: #dc3545;">{wp['weak_phrase']}</span><br>
                            &#9989; <span style="color: #28a745;">{wp['strong_alternative']}</span><br>
                            <small><em>Context: "{wp['context'][:100]}..."</em></small>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.success("✅ No weak phrases detected! Your language is strong.")
            
            with tab2:
                st.markdown("### 💪 Action Verb Suggestions")
                
                if rewrite['action_verbs']:
                    st.info("Diversify your resume with these powerful action verbs")
                    
                    for category, verbs in rewrite['action_verbs'].items():
                        with st.expander(f"📌 {category} Verbs"):
                            st.markdown(", ".join(verbs))
                            st.caption(f"Use these verbs to describe {category.lower()} activities")
                else:
                    st.success("✅ Good variety of action verbs!")
            
            with tab3:
                st.markdown("### 📊 Quantification Opportunities")
                
                if rewrite['quantification']:
                    st.warning(f"Found {len(rewrite['quantification'])} opportunities to add metrics")
                    
                    for idx, opp in enumerate(rewrite['quantification'][:8], 1):
                        st.markdown(f"""
                        <div class="improvement-card">
                            <strong>Opportunity {idx}</strong><br>
                            <em>"{opp['context'][:100]}..."</em><br>
                            &#128161; <strong>{opp['suggestion']}</strong>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("""
                    **💡 Tips for Quantification:**
                    - Add percentages (improved by X%)
                    - Include numbers (managed team of X)
                    - Specify amounts (saved $X)
                    - Show scale (X users, X projects)
                    """)
                else:
                    st.success("✅ Good use of quantifiable metrics!")
            
            with tab4:
                st.markdown("### 🎨 Formatting Improvements")
                
                if rewrite['formatting']:
                    for fmt in rewrite['formatting']:
                        st.markdown(f"• {fmt}")
                else:
                    st.success("✅ Formatting looks good!")
                
                st.markdown("---")
                st.markdown("#### 📋 General Formatting Best Practices")
                st.markdown("""
                - Use consistent bullet points (•)
                - Keep dates in same format (MM/YYYY)
                - Use clear section headers
                - Maintain proper spacing
                - Ensure readability with appropriate line breaks
                - Use bold for emphasis on key achievements
                """)
            
            with tab5:
                if show_examples and rewrite['optimized_sections']:
                    st.markdown("### 📝 Optimized Section Examples")
                    
                    sections = rewrite['optimized_sections']
                    
                    # Professional Summary
                    if 'Professional Summary' in sections:
                        st.markdown("#### 💼 Professional Summary Example")
                        st.markdown(f"""
                        <div class="example-box">
                        {sections['Professional Summary']}
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Experience bullets
                    if 'Experience Bullet Points' in sections:
                        st.markdown("#### 📌 Experience Bullet Point Examples")
                        for bullet in sections['Experience Bullet Points']:
                            st.markdown(f"""
                            <div class="example-box">
                            {bullet}
                            </div>
                            """, unsafe_allow_html=True)
                    
                    # Skills section
                    if 'Skills Section' in sections:
                        st.markdown("#### 🔧 Skills Section Example")
                        st.markdown(f"""
                        <div class="example-box">
                        <pre>{sections['Skills Section']}</pre>
                        </div>
                        """, unsafe_allow_html=True)
            
            # DeepSeek R1 Optimizations Section
            if rewrite.get('deepseek_enabled'):
                st.markdown("---")
                st.markdown("## 🤖 DeepSeek R1 - Personalized Optimization")
                
                # Weak Phrases avec DeepSeek
                if rewrite.get('deepseek_weak_phrases'):
                    st.markdown("### 🔄 Weak Phrases (AI-Detected)")
                    
                    for i, phrase in enumerate(rewrite['deepseek_weak_phrases'][:10], 1):
                        with st.expander(f"❌ Phrase {i}: \"{phrase.get('original', 'N/A')}\"", expanded=(i <= 3)):
                            st.markdown(f"**Why it's weak:** {phrase.get('reason', 'Generic language')}")
                            st.markdown(f"**✅ Stronger alternative:**")
                            st.success(phrase.get('alternative', 'Use more specific action verbs'))
                            
                            priority = phrase.get('priority', 'Medium')
                            if priority == 'High':
                                st.error(f"🔥 Priority: {priority}")
                            elif priority == 'Medium':
                                st.warning(f"⚠️ Priority: {priority}")
                            else:
                                st.info(f"💡 Priority: {priority}")
                
                # Action Verbs avec DeepSeek
                if rewrite.get('deepseek_action_verbs'):
                    st.markdown("### 💪 Action Verbs (Contextualized)")
                    
                    verbs = rewrite['deepseek_action_verbs']
                    verb_categories = list(verbs.keys())[:4]  # Max 4 catégories
                    
                    if verb_categories:
                        tabs = st.tabs(verb_categories)
                        
                        for tab, category in zip(tabs, verb_categories):
                            with tab:
                                verb_list = verbs[category]
                                for verb_data in verb_list[:5]:
                                    if isinstance(verb_data, dict):
                                        st.markdown(f"**{verb_data.get('verb', 'N/A')}**")
                                        st.caption(f"Example: {verb_data.get('example', 'N/A')}")
                                    else:
                                        st.markdown(f"**{verb_data}**")
                                    st.markdown("---")
                
                # Quantification avec DeepSeek
                if rewrite.get('deepseek_quantification'):
                    st.markdown("### 📊 Quantification Opportunities (AI-Identified)")
                    
                    for i, quant in enumerate(rewrite['deepseek_quantification'][:5], 1):
                        with st.expander(f"📈 Opportunity {i}", expanded=(i == 1)):
                            st.markdown("**Original (vague):**")
                            st.info(quant.get('original', 'N/A'))
                            
                            st.markdown("**✅ Quantified version:**")
                            st.success(quant.get('quantified', 'N/A'))
                            
                            st.markdown(f"**Impact:** {quant.get('impact', 'Makes achievement measurable')}")
                
                # Formatting avec DeepSeek
                if rewrite.get('deepseek_formatting'):
                    st.markdown("### 🎨 Formatting Improvements (AI-Detected)")
                    
                    for issue in rewrite['deepseek_formatting']:
                        st.markdown(f"• {issue}")
                
                # Priority Recommendations avec DeepSeek
                if rewrite.get('deepseek_recommendations'):
                    st.markdown("### 🎯 Priority Recommendations (AI-Personalized)")
                    
                    recs = rewrite['deepseek_recommendations']
                    
                    if recs.get('high_priority'):
                        st.markdown("#### 🔥 High Priority")
                        for rec in recs['high_priority']:
                            st.error(f"• {rec}")
                    
                    if recs.get('medium_priority'):
                        st.markdown("#### ⚠️ Medium Priority")
                        for rec in recs['medium_priority']:
                            st.warning(f"• {rec}")
                    
                    if recs.get('low_priority'):
                        st.markdown("#### 💡 Low Priority")
                        for rec in recs['low_priority']:
                            st.info(f"• {rec}")
                
                # Texte brut si mode texte
                if rewrite.get('deepseek_text_mode'):
                    with st.expander("📝 DeepSeek Full Response (Text Mode)", expanded=False):
                        st.text_area("Raw Response", rewrite.get('deepseek_raw_response', ''), height=400, key="deepseek_raw")
            
            # Recommendations
            st.markdown("---")
            st.markdown("### 🎯 Priority Recommendations")
            
            for rec in rewrite['recommendations']:
                priority = rec['priority']
                css_class = f"priority-{priority.lower()}"
                
                st.markdown(f"""
                <div class="{css_class}">
                    <strong>🔥 {priority} Priority - {rec['category']}</strong><br>
                    {rec['recommendation']}<br>
                    <small><em>Impact: {rec['impact']}</em></small>
                </div>
                """, unsafe_allow_html=True)
            
            # Action items
            st.markdown("---")
            st.markdown("### ✅ Next Steps")
            st.markdown("""
            1. Review and implement high-priority suggestions
            2. Replace weak phrases with stronger alternatives
            3. Add quantifiable metrics to your achievements
            4. Diversify your action verbs
            5. Ensure consistent formatting throughout
            6. Tailor your resume for specific job applications
            """)
            
            # Download report
            st.markdown("---")
            st.markdown("### 📥 Generate & Download Optimized Resume")
            
            # Check if DeepSeek is enabled
            if rewrite.get('deepseek_enabled'):
                st.info("💡 Click 'Generate Optimized Resume' to create an editable version based on AI suggestions")
                
                # Button to generate optimized resume
                if st.button("✨ Generate Optimized Resume", type="primary"):
                    with st.spinner("🤖 Generating optimized resume content..."):
                        # Get client and profile
                        client = st.session_state.get('rewriter_deepseek_client')
                        profile = rewrite.get('candidate_profile', {})
                        improvements = {
                            'weak_phrases': rewrite.get('deepseek_weak_phrases', []),
                            'action_verbs': rewrite.get('deepseek_action_verbs', {}),
                            'quantifications': rewrite.get('deepseek_quantification', []),
                            'formatting': rewrite.get('deepseek_formatting', [])
                        }
                        
                        # Generate optimized content
                        optimized_content = generate_optimized_resume_content(
                            client, parsed["text"], profile, improvements
                        )
                        
                        # Store in session state
                        st.session_state['optimized_resume_content'] = optimized_content
                        st.success("✅ Optimized resume generated! You can now edit it below.")
                        st.rerun()
                
                # Show editing interface if content is generated
                if 'optimized_resume_content' in st.session_state:
                    st.markdown("---")
                    st.markdown("### ✏️ Edit Your Optimized Resume")
                    st.info("📝 Review and modify the content below before downloading")
                    
                    content = st.session_state['optimized_resume_content']
                    
                    # Editable sections
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        st.markdown("#### Professional Summary")
                        edited_summary = st.text_area(
                            "Summary",
                            value=content.get('summary', ''),
                            height=100,
                            key="edit_summary",
                            label_visibility="collapsed"
                        )
                        
                        st.markdown("#### Professional Experience")
                        # Handle experience as list or string
                        exp_value = content.get('experience', [])
                        if isinstance(exp_value, list):
                            exp_text = '\n'.join(exp_value)
                        else:
                            exp_text = exp_value
                        
                        edited_experience = st.text_area(
                            "Experience",
                            value=exp_text,
                            height=200,
                            key="edit_experience",
                            label_visibility="collapsed"
                        )
                        
                        st.markdown("#### Skills")
                        edited_skills = st.text_area(
                            "Skills",
                            value=content.get('skills', ''),
                            height=100,
                            key="edit_skills",
                            label_visibility="collapsed"
                        )
                        
                        st.markdown("#### Education")
                        edited_education = st.text_area(
                            "Education",
                            value=content.get('education', ''),
                            height=80,
                            key="edit_education",
                            label_visibility="collapsed"
                        )
                    
                    with col2:
                        st.markdown("#### Actions")
                        st.markdown("**Preview Tips:**")
                        st.caption("• Review each section carefully")
                        st.caption("• Check for accuracy")
                        st.caption("• Ensure metrics are realistic")
                        st.caption("• Verify all information")
                        
                        st.markdown("---")
                        
                        # Save edited content
                        if st.button("💾 Save Edits", use_container_width=True):
                            st.session_state['optimized_resume_content'] = {
                                'summary': edited_summary,
                                'experience': edited_experience,
                                'skills': edited_skills,
                                'education': edited_education
                            }
                            st.success("✅ Edits saved!")
                        
                        # Reset button
                        if st.button("🔄 Reset to AI Version", use_container_width=True):
                            # Regenerate content
                            with st.spinner("Regenerating..."):
                                client = st.session_state.get('rewriter_deepseek_client')
                                profile = rewrite.get('candidate_profile', {})
                                improvements = {
                                    'weak_phrases': rewrite.get('deepseek_weak_phrases', []),
                                    'action_verbs': rewrite.get('deepseek_action_verbs', {}),
                                    'quantifications': rewrite.get('deepseek_quantification', []),
                                    'formatting': rewrite.get('deepseek_formatting', [])
                                }
                                optimized_content = generate_optimized_resume_content(
                                    client, parsed["text"], profile, improvements
                                )
                                st.session_state['optimized_resume_content'] = optimized_content
                                st.success("✅ Reset to AI version!")
                                st.rerun()
                    
                    # Download buttons
                    st.markdown("---")
                    st.markdown("### 📥 Download Your Resume")
                    
                    # Prepare optimized data for download
                    download_data = rewrite.copy()
                    download_data['deepseek_optimized'] = {
                        'summary': edited_summary,
                        'experience': [line for line in edited_experience.split('\n') if line.strip()],
                        'skills': edited_skills,
                        'education': edited_education
                    }
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # Génération DOCX
                        try:
                            with st.spinner("Generating DOCX..."):
                                docx_buffer = generate_optimized_resume_docx(parsed["text"], download_data)
                            
                            st.download_button(
                                label="📄 Download as DOCX",
                                data=docx_buffer,
                                file_name=f"optimized_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                help="Download optimized resume in Word format",
                                use_container_width=True
                            )
                        except Exception as e:
                            st.error(f"Error generating DOCX: {str(e)}")
                    
                    with col2:
                        # Génération PDF
                        try:
                            with st.spinner("Generating PDF..."):
                                pdf_buffer = generate_optimized_resume_pdf(parsed["text"], download_data)
                            
                            st.download_button(
                                label="📕 Download as PDF",
                                data=pdf_buffer,
                                file_name=f"optimized_resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                mime="application/pdf",
                                help="Download optimized resume in PDF format",
                                use_container_width=True
                            )
                        except ImportError:
                            st.warning("⚠️ PDF generation requires reportlab. Install with: pip install reportlab")
                        except Exception as e:
                            st.error(f"Error generating PDF: {str(e)}")
            else:
                st.info("💡 Enable DeepSeek R1 to generate a complete optimized version of your resume downloadable in PDF/DOCX")
            
            # Standard report downloads
            st.markdown("---")
            st.markdown("### 📥 Download Rewriting Report")
            
            # Generate comprehensive report
            report_text = f"""
UtopiaHire - Resume Rewriting Suggestions Report
Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
================================================

IMPROVEMENT POTENTIAL: {rewrite.get('improvement_score', 'N/A')}/100

WEAK PHRASES TO REPLACE ({len(rewrite.get('weak_phrases', []))})
========================
"""
            
            for idx, wp in enumerate(rewrite.get('weak_phrases', []), 1):
                report_text += f"""
{idx}. Weak: "{wp['weak_phrase']}"
   Strong: "{wp['strong_alternative']}"
   Context: "{wp['context'][:80]}..."
"""
            
            report_text += f"""

ACTION VERBS BY CATEGORY
========================
"""
            for category, verbs in rewrite.get('action_verbs', {}).items():
                report_text += f"\n{category}:\n  {', '.join(verbs)}\n"
            
            report_text += f"""

QUANTIFICATION OPPORTUNITIES ({len(rewrite.get('quantification', []))})
============================
"""
            for idx, opp in enumerate(rewrite.get('quantification', []), 1):
                report_text += f"""
{idx}. Context: "{opp['context'][:80]}..."
   Suggestion: {opp['suggestion']}
"""
            
            report_text += f"""

FORMATTING IMPROVEMENTS
=======================
"""
            for fmt in rewrite.get('formatting', []):
                report_text += f"• {fmt}\n"
            
            report_text += f"""

CONTENT ENHANCEMENTS
====================
"""
            for enh in rewrite.get('content_enhancements', []):
                report_text += f"• {enh}\n"
            
            report_text += f"""

PRIORITY RECOMMENDATIONS
========================
"""
            for rec in rewrite.get('recommendations', []):
                report_text += f"""
[{rec['priority']} Priority] {rec['category']}
  Recommendation: {rec['recommendation']}
  Impact: {rec['impact']}
"""
            
            report_text += f"""

OPTIMIZED SECTION EXAMPLES
==========================
"""
            for section_name, content in rewrite.get('optimized_sections', {}).items():
                report_text += f"\n{section_name}:\n"
                if isinstance(content, list):
                    for item in content:
                        report_text += f"  {item}\n"
                else:
                    report_text += f"  {content}\n"
            
            report_text += """

================================================
Report generated by UtopiaHire - AI Career Architect
Next Steps: Implement suggestions and re-analyze your resume
"""
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Text report download
                st.download_button(
                    label="📄 Download Text Report",
                    data=report_text,
                    file_name=f"resume_rewriting_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    help="Download complete rewriting suggestions as text file"
                )
            
            with col2:
                # JSON report download
                json_report = json.dumps(rewrite, indent=2, default=str)
                st.download_button(
                    label="📊 Download JSON Report",
                    data=json_report,
                    file_name=f"resume_rewriting_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    help="Download structured data for programmatic use"
                )

if __name__ == "__main__":
    main()
